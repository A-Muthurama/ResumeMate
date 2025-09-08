# app.py (final - ReportLab only)
import os
import io
import json
import re
import tempfile
from flask import Flask, render_template, request, flash, send_file, redirect, url_for
import google.generativeai as genai
from docx import Document
import fitz  # PyMuPDF
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4, letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# ---------------- CONFIG ----------------
load_dotenv()
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET", "dev-secret")

# Gemini API key
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
else:
    genai = None

GEMINI_MODEL_NAME = "gemini-2.5-flash"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
JOBS_PATH = os.path.join(BASE_DIR, "jobs.json")

# load jobs
if os.path.exists(JOBS_PATH):
    with open(JOBS_PATH, "r", encoding="utf-8") as f:
        JOBS_DB = json.load(f)
else:
    JOBS_DB = []

# Predefined skill set (fallback local scorer)
KNOWN_SKILLS = {
    "python", "java", "c++", "javascript", "html", "css", "sql",
    "machine learning", "deep learning", "nlp", "data analysis",
    "excel", "power bi", "tableau", "cloud", "aws", "azure",
    "flask", "django", "react", "node.js", "git", "docker", "kubernetes"
}

# ---------------- Helpers ----------------
def ask_gemini_json(prompt: str):
    if not genai:
        return None
    try:
        model = genai.GenerativeModel(GEMINI_MODEL_NAME)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        app.logger.error("Gemini error (json): %s", e)
        return None

def ask_gemini_text(prompt: str):
    if not genai:
        return None
    try:
        model = genai.GenerativeModel(GEMINI_MODEL_NAME)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        app.logger.error("Gemini error (text): %s", e)
        return None

def extract_skills_from_text(text: str):
    text = (text or "").lower()
    found = set()
    for skill in KNOWN_SKILLS:
        if re.search(rf"\b{re.escape(skill)}\b", text):
            found.add(skill)
    return found

def ats_score_local(resume_text, job_desc):
    resume_skills = extract_skills_from_text(resume_text)
    job_skills = extract_skills_from_text(job_desc)

    matched = resume_skills.intersection(job_skills)
    missing = job_skills - resume_skills

    score = int((len(matched) / len(job_skills) * 100)) if job_skills else 0
    return {
        "ats_score": score,
        "matched_skills": sorted(list(matched)),
        "missing_skills": sorted(list(missing)),
        "suggestions": (
            "Add missing skills to improve your ATS score. "
            f"Missing: {', '.join(sorted(list(missing))) if missing else 'None'}"
        )
    }

def safe_extract_text(file_storage):
    if not file_storage:
        return ""
    try:
        from utils.parser import extract_text as external_extract
    except Exception:
        external_extract = None
    if external_extract:
        try:
            file_storage.seek(0)
            txt = external_extract(file_storage)
            if isinstance(txt, str) and txt.strip():
                return txt
        except Exception as e:
            app.logger.debug("external_extract failed: %s", e)
    filename = (file_storage.filename or "").lower()
    file_storage.seek(0)
    data = file_storage.read()
    if not data:
        return ""
    if filename.endswith(".pdf") or (data[:4] == b"%PDF"):
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            return "\n".join([page.get_text() for page in doc]).strip()
        except Exception as e:
            app.logger.error("PyMuPDF parsing failed: %s", e)
            return ""
    if filename.endswith(".docx"):
        try:
            import io
            doc = Document(io.BytesIO(data))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()]).strip()
        except Exception as e:
            app.logger.error("docx parsing failed: %s", e)
            return ""
    try:
        return data.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""

# ---------------- Routes ----------------
@app.route("/")
def home():
    return render_template("index.html")

@app.route("/resume", methods=["GET", "POST"])
def resume_creation():
    if request.method == "POST":
        data = {
            "name": request.form.get("name", ""),
            "email": request.form.get("email", ""),
            "phone": request.form.get("phone", ""),
            "linkedin": request.form.get("linkedin", ""),
            "summary": request.form.get("summary", ""),
            "skills": request.form.get("skills", ""),
            "education": request.form.get("education", ""),
            "projects": request.form.get("projects", ""),
            "achievements": request.form.get("achievements", ""),
            "certifications": request.form.get("certifications", ""),
            "extras": request.form.get("extras", "")
        }
        pdf_url = url_for("download_pdf", **data)
        docx_url = url_for("download_docx", **data)
        return render_template("resume.html", resume_text=data, pdf_url=pdf_url, docx_url=docx_url)
    return render_template("resume.html")

# ✅ PDF Download using ReportLab
@app.route("/download_pdf")
def download_pdf():
    data = request.args.to_dict()
    pdf_path = os.path.join(tempfile.gettempdir(), "resume.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, data.get("name", "Your Name"))
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 70, f"Email: {data.get('email','')} | Phone: {data.get('phone','')} | LinkedIn: {data.get('linkedin','')}")
    y = height - 100
    for section, content in data.items():
        if section in ["name", "email", "phone", "linkedin"]:
            continue
        if content.strip():
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, section.capitalize())
            y -= 20
            c.setFont("Helvetica", 10)
            for line in content.split("\n"):
                c.drawString(70, y, line.strip())
                y -= 15
            y -= 10
    c.save()
    return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

# ✅ DOCX Download
@app.route("/download_docx")
def download_docx():
    data = request.args.to_dict()
    doc = Document()
    doc.add_heading(data.get("name", "Your Name"), 0)
    doc.add_paragraph(f"Email: {data.get('email','')} | Phone: {data.get('phone','')} | LinkedIn: {data.get('linkedin','')}")
    for section, content in data.items():
        if section in ["name", "email", "phone", "linkedin"]:
            continue
        if content.strip():
            doc.add_heading(section.capitalize(), level=1)
            doc.add_paragraph(content)
    doc_path = os.path.join(tempfile.gettempdir(), "resume.docx")
    doc.save(doc_path)
    return send_file(doc_path, as_attachment=True, download_name="resume.docx")

# Job Matching
@app.route("/match", methods=["GET", "POST"])
def job_matching():
    matches = None
    if request.method == "POST":
        uploaded = request.files.get("resume_file")
        resume_text = request.form.get("resume_text", "").strip()
        if uploaded and uploaded.filename:
            resume_text = safe_extract_text(uploaded)
        if not resume_text:
            flash("Please upload or paste a resume.", "warning")
            return render_template("match.html")
        if genai:
            job_list_str = "\n".join([f"- {j.get('title','')}: {', '.join(j.get('keywords', []))}" for j in JOBS_DB])
            prompt = (
                "You are a job-matching assistant. Given a resume and job roles, return JSON array of "
                "{title, score (0-100), matched_skills:[], missing_skills:[]}.\n\n"
                f"Jobs:\n{job_list_str}\n\nResume:\n{resume_text}\n\nReturn JSON only."
            )
            gem_text = ask_gemini_json(prompt)
            try:
                m = re.search(r"(\[.*\])", gem_text, re.S)
                parsed = json.loads(m.group(1) if m else gem_text)
                matches = parsed
            except Exception:
                matches = []
        else:
            matches = []
        matches = sorted(matches, key=lambda x: x.get("score", 0), reverse=True)
    return render_template("match.html", matches=matches)

# ATS Scoring
@app.route("/score", methods=["GET", "POST"])
def ats_scoring():
    result = None
    if request.method == "POST":
        uploaded = request.files.get("resume_file")
        resume_text = request.form.get("resume_text", "").strip()
        job_desc = request.form.get("job_desc", "").strip()
        if uploaded and uploaded.filename:
            resume_text = safe_extract_text(uploaded)
        if not resume_text or not job_desc:
            flash("Upload/paste resume and provide job description.", "warning")
            return render_template("score.html")
        if genai:
            prompt = (
                "You are an ATS-scoring assistant. Compare resume and job description and return JSON: "
                "{ats_score:int, matched_skills:[], missing_skills:[], suggestions:'...'}\n\n"
                f"Resume:\n{resume_text}\n\nJob Description:\n{job_desc}\n\nReturn JSON only."
            )
            gem_text = ask_gemini_json(prompt)
            try:
                m = re.search(r"(\{.*\})", gem_text, re.S)
                parsed = json.loads(m.group(1) if m else gem_text)
                result = parsed
            except Exception:
                result = ats_score_local(resume_text, job_desc)
        else:
            result = ats_score_local(resume_text, job_desc)
    return render_template("score.html", result=result)

@app.route("/healthz")
def healthz():
    return "OK", 200

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
